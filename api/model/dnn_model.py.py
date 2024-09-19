import sys
import pandas as pd
import numpy as np
import tensorflow as tf
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, Input
from scikeras.wrappers import KerasClassifier
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, roc_curve, auc
from imblearn.over_sampling import SMOTE
from imblearn.pipeline import make_pipeline
import matplotlib.pyplot as plt
import random
import io
import traceback
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from io import BytesIO
from tensorflow.keras.optimizers import Nadam
from tensorflow.keras.optimizers import RMSprop


def add_dataframe_to_doc(doc, dataframe, title):
    doc.add_heading(title, level=1)
    if not dataframe.empty:
        # Create table with add headings
        table = doc.add_table(rows=(dataframe.shape[0] + 1), cols=dataframe.shape[1])
        table.style = 'Table Grid'
        
        # Headings configuration of the table
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(dataframe.columns):
            hdr_cells[i].text = str(col_name)
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(11)  # Establecer el tamaño de fuente a 11
        
        # Fill the table with the data
        for index, row in enumerate(dataframe.itertuples(), start=1):
            row_cells = table.rows[index].cells
            for j, value in enumerate(row[1:]):  
                if isinstance(value, float):
                    rounded_value = round(value, 3)  #Rounded to 3 decimal places
                else:
                    rounded_value = value
                row_cells[j].text = str(rounded_value)
                row_cells[j].paragraphs[0].runs[0].font.size = Pt(11)  # Font size 11

        # Adjust the column width according to page size 
        table_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        col_width = table_width / len(dataframe.columns)
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width

    else:
        doc.add_paragraph("No data available.")

def main():
    try:
        print("Python script started", file=sys.stderr)

        # Set seeds for reproducibility
        np.random.seed(42)
        random.seed(42)
        tf.random.set_seed(42)

        # Read the Excel file from stdin
        print("Reading Excel file from stdin", file=sys.stderr)
        excel_data = sys.stdin.buffer.read()
        financial_data = pd.read_excel(io.BytesIO(excel_data))
        financial_ratios_data = pd.read_excel(io.BytesIO(excel_data), sheet_name='ratios')
        financial_data.fillna(0, inplace=True)
        financial_data['Fecha'] = financial_data['Fecha'].astype(str)

        print("Excel file read successfully", file=sys.stderr)

        # Convert all categorical columns to string and compute medians
        categorical_features = ['Criterio estado financiero', 'Cuenta', 'Detalle', 'Item', 'Fecha']
        for col in categorical_features:
            financial_data[col] = financial_data[col].astype(str)
        median_values = financial_data.groupby('Criterio estado financiero')['Monto'].median()
        financial_data['Label'] = financial_data.apply(lambda x: 1 if x['Monto'] > median_values[x['Criterio estado financiero']] else 0, axis=1)

        # Calculate financial ratios
        total_sales = financial_ratios_data.loc[financial_ratios_data['Criterio Resultado'] == 'VENTAS', 'Valor Resultado'].sum()
        cost_of_sales = financial_ratios_data.loc[financial_ratios_data['Criterio Resultado'] == 'COSTO DE VENTAS', 'Valor Resultado'].sum()
        gross_profit = total_sales - cost_of_sales
        total_assets = financial_data.loc[financial_data['Criterio estado financiero'] == 'ACTIVO FIJO NETO', 'Monto'].sum()
        equity = financial_data.loc[financial_data['Criterio estado financiero'] == 'PATRIMONIO', 'Monto'].sum()
        total_liabilities = financial_data.loc[financial_data['Criterio estado financiero'].str.contains('PASIVO'), 'Monto'].sum()
        current_liabilities = financial_data.loc[financial_data['Cuenta'] == 'PASIVO CORRIENTE', 'Monto'].sum()
        cash = financial_data.loc[financial_data['Cuenta'].str.contains('Efectivo y Equivalentes de Efectivo'), 'Monto'].sum()

        financial_data['Margen de Ganancia Bruta sobre Ventas'] = gross_profit / total_sales
        financial_data['Rentabilidad de los Activos Totales'] = gross_profit / total_assets
        financial_data['Rentabilidad del Patrimonio antes de Impuestos'] = (gross_profit - cost_of_sales) / equity
        financial_data['Rentabilidad sobre el Patrimonio'] = gross_profit / equity
        financial_data['Cobertura de Efectivo sobre Pasivos Totales'] = cash / total_liabilities
        financial_data['Ratio de Endeudamiento'] = total_liabilities / equity

        X = financial_data.drop(['Monto', 'Label'], axis=1)
        y = financial_data['Label']
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

        preprocessor = ColumnTransformer(transformers=[
            ('cat', OneHotEncoder(handle_unknown='ignore'), categorical_features)
        ], remainder='passthrough')

        def build_model(input_shape):
            input_layer = Input(shape=(input_shape,))
            dense1 = Dense(128, activation='relu')(input_layer)
            dense2 = Dense(128, activation='relu')(dense1)
            dense3 = Dense(128, activation='relu')(dense2)
            output = Dense(1, activation='sigmoid')(dense3)
            model = Model(inputs=input_layer, outputs=output)
            optimizer = RMSprop(learning_rate=0.001, rho=0.9)  
            model.compile(optimizer=optimizer, loss='binary_crossentropy', metrics=['accuracy', 'precision', 'recall', 'auc'])
            return model

        smote_pipeline = make_pipeline(
            preprocessor,
            SMOTE(random_state=42),
            KerasClassifier(model=build_model, epochs=250, batch_size=32, verbose=1)
        )

        input_shape = preprocessor.fit_transform(X_train).shape[1]
        smote_pipeline.set_params(kerasclassifier__model__input_shape=input_shape)
        smote_pipeline.fit(X_train, y_train)

        predictions = smote_pipeline.predict(X_test)
        probabilities = smote_pipeline.predict_proba(X_test)[:, 1]
        fpr, tpr, thresholds = roc_curve(y_test, probabilities)
        roc_auc = auc(fpr, tpr)

        plt.figure()
        plt.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (area = {roc_auc:.2f})')
        plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
        plt.xlabel('False Positive Rate')
        plt.ylabel('True Positive Rate')
        plt.title('Receiver Operating Characteristic')
        plt.legend(loc='lower right')
        plt.savefig("roc_curve.png")
        plt.close()

        anomaly_threshold = 0.90
        anomalies_indices = np.where((probabilities > anomaly_threshold) & (predictions == 1))[0]
        anomalies_indices = X_test.iloc[anomalies_indices].index 

        thresholds = {
            'Margen de Ganancia Bruta sobre Ventas': 1.0,  
            'Rentabilidad de los Activos Totales': 2.0, 
            'Rentabilidad del Patrimonio antes de Impuestos': 7.0, 
            'Rentabilidad sobre el Patrimonio': 4.0,
            'Cobertura de Efectivo sobre Pasivos Totales': 5.0, 
            'Ratio de Endeudamiento': 8.0,
        }

        additional_ratios_indices = []
        for ratio, threshold in thresholds.items():
            if ratio in X_test.columns:
                idx = X_test[X_test[ratio] > threshold].index
                additional_ratios_indices.extend(idx)

        combined_indices = np.union1d(anomalies_indices, additional_ratios_indices)

        doc = Document()
        doc.add_heading('Informe de detección de anomalías financieras', 0)

        doc.add_heading('Métricas de rendimiento - IA ', level=1)
        doc.add_paragraph("Las siguientes métricas representan el rendimiento del modelo de IA utilizado en la detección de anomalías financieras.")
        doc.add_paragraph("Definiciones:")
        doc.add_paragraph("Accuracy: Proporción de predicciones correctas.")
        doc.add_paragraph("Precision: Proporción de predicciones positivas correctas.")
        doc.add_paragraph("Recall: Proporción de verdaderos positivos detectados (anomalías).")
        doc.add_paragraph("F1 Score: Media armónica de precision y recall.")
        # Utilizar una tabla para mostrar las métricas
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        metrics = [
             ("Accuracy", f"{accuracy_score(y_test, predictions):.2f}"),
             ("Precision", f"{precision_score(y_test, predictions):.2f}"),
             ("Recall", f"{recall_score(y_test, predictions):.2f}"),
             ("F1 Score", f"{f1_score(y_test, predictions):.2f}"),
             ("ROC AUC", f"{roc_auc:.2f}")
        ]
        
        for metric, value in metrics:
           row_cells = table.add_row().cells
           row_cells[0].text = metric
           row_cells[1].text = value
                          

        #doc.add_heading('Curva de Característica Operativa del Receptor - IA', level=1)
        #doc.add_picture('roc_curve.png', width=Inches(5))
        #doc.add_paragraph("La curva ROC es una representación gráfica de la sensibilidad frente a la especificidad para un sistema clasificador binario en función de diferentes umbrales de decisión. Cuanto mayor sea el área bajo la curva (AUC), mejor será el rendimiento del modelo.")

        if len(combined_indices) > 0:
            anomalies = X_test.loc[combined_indices]
            warning_title = doc.add_heading("Observaciones")
            warning_paragraph = doc.add_paragraph("El modelo ha identificado ciertas anomalías que requieren una investigación más profunda para entender sus causas y posibles implicaciones. Recomendamos que el personal auditor o la persona encargada realice una revisión detallada de los documentos y transacciones relacionadas para validar la exactitud de los datos y ajustar los registros financieros según sea necesario.")
            warning_paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            warning_title.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            add_dataframe_to_doc(doc, anomalies, "Posibles anomalías detectadas")
        else:
            doc.add_paragraph('Ninguna anomalía fue detectada.')
            
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        sys.stdout.buffer.write(buffer.getvalue()) #reemplaza doc.save()

        print("Report generated successfully: financial_anomalies_report.docx", file=sys.stderr)

    except Exception as e:
        print(f"An error occurred: {str(e)}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
